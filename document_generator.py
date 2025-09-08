from __future__ import annotations
import os
import re
import sys
import copy

from docx import Document
from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime
from typing import List, Dict, Any, Tuple

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

from num2words import num2words
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------- CONSTANTS (edit) ----------------------
SERVICE_ACCOUNT_KEYFILE = "asset-acts-15fc23ac9cd4.json"

ASSETS_SPREADSHEET_ID = "1XbJuTGtwiQNFzxqOM5F1ehOE5PRmd1GFzi2HFmZnoWE"
ASSETS_SHEET_NAME = "list"

DEPARTMENTS_SPREADSHEET_ID = "1N45PcHU-YgpcYEyXDbQRCG94I-BWbrultIyqRr5Z4N8"
DEPARTMENTS_SHEET_NAME = "Department"

TEMPLATE_DOC_ID = "1s6ZWVpMY3g3LU7IYM4fFyRG6mPfh0yy5_xzq5KVMIsI"
OUTPUT_FOLDER_ID = "1DHm3RDZU_0-dKKKqpw5qHPwV8E7MIfA4"  # Drive folder id (optional if local fallback)

OUTPUT_LOCAL_DIR = "output_docs"  # where .docx files will be saved if Drive unavailable

# Columns (1-based)
COL_ID = 1
COL_NAME = 2
COL_INVENTORY_NUMBER = 3
COL_UNIT = 4
COL_QUANTITY = 5
COL_PRICE = 6
COL_OWNERS = 7
COL_DATE = 8
COL_GENERATE_FLAG = 9

DEPT_COL_CODE = 1
DEPT_COL_TYPE = 2
DEPT_COL_STATUS = 3
DEPT_COL_POSITION = 4
DEPT_COL_FULLNAME = 5
DEPT_COL_NORMALIZED = 6

DATE_OUTPUT_FORMAT = "%Y%m%d"
FILE_NAME_PATTERN = "Акт. {deptname}"  
ROUND_DECIMALS = 2
THOUSAND_SEPARATOR = " "
DECIMAL_SEPARATOR = ","
CURRENCY_SUFFIX = ""
ALLOW_ROUNDING_ADJUST = True
TABLE_PLACEHOLDER_TOKEN = "%TABLE_PLACEHOLDER%"

# ----------------------------------------------------------------

LOG_DIR = "logs"
os.makedirs(LOG_DIR, exist_ok=True)
_logfile_path = os.path.join(LOG_DIR, f"asset_acts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
_logfile = open(_logfile_path, "a", encoding="utf-8")


def _write(level: str, msg: str) -> None:
    line = f"{level.upper()}: {msg}"
    print(line, flush=True)
    _logfile.write(line + "\n")
    _logfile.flush()


def info(msg: str) -> None:
    _write("INFO", msg)


def warning(msg: str) -> None:
    _write("WARNING", msg)


def error(msg: str) -> None:
    _write("ERROR", msg)


def summary(msg: str) -> None:
    _write("SUMMARY", msg)


def check_constants() -> None:
    missing = []
    if not os.path.isfile(SERVICE_ACCOUNT_KEYFILE):
        missing.append(f"SERVICE_ACCOUNT_KEYFILE file not found: {SERVICE_ACCOUNT_KEYFILE}")
    for keyname, value in (
        ("ASSETS_SPREADSHEET_ID", ASSETS_SPREADSHEET_ID),
        ("DEPARTMENTS_SPREADSHEET_ID", DEPARTMENTS_SPREADSHEET_ID),
        ("TEMPLATE_DOC_ID", TEMPLATE_DOC_ID),
    ):
        if not value or value.startswith("FOLDER_ID_HERE") and keyname == "ASSETS_SPREADSHEET_ID":
            # treat folder optional because local fallback exists
            pass
    if missing:
        for m in missing:
            error(m)
        raise SystemExit("Missing critical constants; aborting.")


SCOPES = [
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/documents",
    "https://www.googleapis.com/auth/spreadsheets.readonly",
]


def build_services():
    creds = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_KEYFILE, scopes=SCOPES)
    sheets = build("sheets", "v4", credentials=creds, cache_discovery=False)
    drive = build("drive", "v3", credentials=creds, cache_discovery=False)
    docs = build("docs", "v1", credentials=creds, cache_discovery=False)
    return sheets, drive, docs

def ensure_file_is_spreadsheet(drive_service, file_id: str, label: str) -> None:
    """
    Verify that `file_id` exists and is a Google Spreadsheet.
    If not: log ERROR and exit.
    """
    try:
        meta = drive_service.files().get(fileId=file_id, fields="id, name, mimeType").execute()
    except HttpError as e:
        error(f"Cannot fetch {label} (id={file_id}): {e}")
        raise SystemExit(1)
    mime = meta.get("mimeType", "")
    if mime != "application/vnd.google-apps.spreadsheet":
        error(f"{label} (id={file_id}) is not a Google Spreadsheet (mimeType={mime}). Please provide the correct spreadsheet ID.")
        raise SystemExit(1)
    info(f"{label} found: {meta.get('name', '<untitled>')} (id={file_id})")


def parse_number(s: str) -> Decimal:
    if s is None:
        raise ValueError("Empty numeric")
    st = str(s).strip()
    st = st.replace("\xa0", "").replace(" ", "")
    st = st.replace(",", ".")
    if st == "":
        raise ValueError("Empty numeric")
    return Decimal(st)


def fmt_number(val: Decimal) -> str:
    q = val.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    s = f"{q:,.2f}"
    s = s.replace(",", "TEMP_THOUS").replace(".", "TEMP_DEC")
    s = s.replace("TEMP_THOUS", THOUSAND_SEPARATOR).replace("TEMP_DEC", DECIMAL_SEPARATOR)
    return s + (CURRENCY_SUFFIX or "")


def is_true_flag(s: str) -> bool:
    if s is None:
        return False
    return str(s).strip().lower() in ("true", "1", "yes", "y", "так")


def normalize_code(token: str) -> str:
    return re.sub(r"\s+", "", token).upper()


def read_sheet_values(sheets_service, spreadsheet_id: str, sheet_name: str):
    rng = f"{sheet_name}"
    try:
        res = sheets_service.spreadsheets().values().get(spreadsheetId=spreadsheet_id, range=rng).execute()
        return res.get("values", [])
    except HttpError as e:
        msg = str(e)
        if "This operation is not supported for this document" in msg or "not supported for this document" in msg:
            error(f"Sheets API: file id {spreadsheet_id} is not a spreadsheet or not supported by Sheets API. Check that the ID points to a Google Sheet and the service account has access.")
            raise SystemExit(1)
        error(f"Sheets API error for spreadsheet={spreadsheet_id}, range={rng}: {e}")
        raise


def load_departments(sheets_service) -> Dict[str, Dict[str, str]]:
    vals = read_sheet_values(sheets_service, DEPARTMENTS_SPREADSHEET_ID, DEPARTMENTS_SHEET_NAME)
    depts = {}
    if not vals or len(vals) < 2:
        warning("Departments sheet empty or missing rows.")
        return depts
    for i, row in enumerate(vals[1:], start=2):
        code = (row[DEPT_COL_CODE - 1] if len(row) >= DEPT_COL_CODE else "").strip()
        if not code:
            warning(f"Departments row {i} missing code; skipping.")
            continue
        key = normalize_code(code)
        depts[key] = {
            "code": row[DEPT_COL_CODE - 1] if len(row) >= DEPT_COL_CODE else "",
            "type": row[DEPT_COL_TYPE - 1] if len(row) >= DEPT_COL_TYPE else "",
            "status": row[DEPT_COL_STATUS - 1] if len(row) >= DEPT_COL_STATUS else "",
            "position": row[DEPT_COL_POSITION - 1] if len(row) >= DEPT_COL_POSITION else "",
            "fullname": row[DEPT_COL_FULLNAME - 1] if len(row) >= DEPT_COL_FULLNAME else "",
            "normalized": row[DEPT_COL_NORMALIZED - 1] if len(row) >= DEPT_COL_NORMALIZED else "",
        }        
    return depts


def parse_owner_token(tok: str) -> Tuple[str, int, bool]:
    tok = tok.strip()
    m = re.match(r"^(.*?)-\s*([0-9]+)\s*$", tok)
    if m:
        return m.group(1).strip(), int(m.group(2)), True
    return tok, None, False


def parse_assets(sheets_service, departments: Dict[str, Dict[str, str]]):
    vals = read_sheet_values(sheets_service, ASSETS_SPREADSHEET_ID, ASSETS_SHEET_NAME)
    if not vals or len(vals) < 2:
        info("No assets rows found.")
        return {}, {"rows_processed": 0, "rows_skipped": 0, "owners_skipped": 0, "total_items_in_acts": 0, "total_value_generated": Decimal("0.00")}

    rows_processed = 0
    rows_skipped = 0
    owners_skipped = 0
    per_owner: Dict[str, Dict[str, Any]] = {}
    total_items_in_acts = 0
    total_value_generated = Decimal("0.00")

    for rindex, row in enumerate(vals[1:], start=2):
        if not any(cell.strip() for cell in row if isinstance(cell, str)):
            continue
        gen_flag = row[COL_GENERATE_FLAG - 1] if len(row) >= COL_GENERATE_FLAG else ""
        if not is_true_flag(gen_flag):
            rows_skipped += 1
            continue
        try:
            name = row[COL_NAME - 1] if len(row) >= COL_NAME else ""
            invnum = row[COL_INVENTORY_NUMBER - 1] if len(row) >= COL_INVENTORY_NUMBER else ""
            unit = row[COL_UNIT - 1] if len(row) >= COL_UNIT else ""
            qty_raw = row[COL_QUANTITY - 1] if len(row) >= COL_QUANTITY else ""
            price_raw = row[COL_PRICE - 1] if len(row) >= COL_PRICE else ""
            owners_raw = row[COL_OWNERS - 1] if len(row) >= COL_OWNERS else ""
            if not name:
                warning(f"Row {rindex} missing name; skipping.")
                rows_skipped += 1
                continue
            qty = int(parse_number(qty_raw))
            price = parse_number(price_raw)
            if qty <= 0:
                error(f"Row {rindex} has non-positive quantity {qty}; skip row.")
                rows_skipped += 1
                continue
            unit_price = (price / Decimal(qty)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        except Exception as e:
            error(f"Row {rindex} parse error: {e}; skipping row.")
            rows_skipped += 1
            continue

        tokens = [t.strip() for t in str(owners_raw).split(",") if t.strip()]
        if not tokens:
            error(f"Row {rindex} no owners; skip.")
            rows_skipped += 1
            continue

        token_infos = []
        any_explicit = False
        for tok in tokens:
            base, num, explicit = parse_owner_token(tok)
            if explicit:
                any_explicit = True
            token_infos.append((base, num, explicit))

        if any_explicit:
            if not all(t[2] for t in token_infos):
                error(f"Row {rindex} mixed explicit and implicit owners; skip.")
                rows_skipped += 1
                continue
            total_spec = sum(t[1] for t in token_infos)
            if total_spec != qty:
                error(f"Row {rindex} owner counts sum {total_spec} != quantity {qty}; skip.")
                rows_skipped += 1
                continue
        else:
            if len(token_infos) != 1:
                error(f"Row {rindex} ambiguous multiple owners without counts; skip.")
                rows_skipped += 1
                continue
            base = token_infos[0][0]
            token_infos[0] = (base, qty, True)

        owners_for_row = []
        for base, num, _ in token_infos:
            key = normalize_code(base)
            dept = departments.get(key)
            if not dept:
                error(f"Row {rindex} owner '{base}' not found; skipping this owner entry.")
                owners_skipped += 1
                continue
            owners_for_row.append((key, int(num), dept))

        if not owners_for_row:
            info(f"Row {rindex}: all owners skipped; skipping row.")
            rows_skipped += 1
            continue

        owner_sums = []
        for (_, oqty, _) in owners_for_row:
            owner_sums.append((unit_price * Decimal(oqty)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))
        sum_owner_sums = sum(owner_sums)
        price_q = price.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        if sum_owner_sums != price_q:
            diff = price_q - sum_owner_sums
            if ALLOW_ROUNDING_ADJUST:
                owner_sums[-1] = (owner_sums[-1] + diff).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
                warning(f"Row {rindex} rounding adjusted by {diff} on last owner.")
            else:
                warning(f"Row {rindex} rounding mismatch {price_q - sum_owner_sums}; continuing.")

        for (key, oqty, dept), osum in zip(owners_for_row, owner_sums):
            if key not in per_owner:
                per_owner[key] = {"dept": dept, "items": [], "tot_qty": 0, "tot_sum": Decimal("0.00")}
            per_owner[key]["items"].append({
                "name": name,
                "inventory": invnum,
                "unit": unit,
                "qty": int(oqty),
                "unit_price": unit_price,
                "sum": osum,
                "note": "",
            })
            per_owner[key]["tot_qty"] += int(oqty)
            per_owner[key]["tot_sum"] += osum
            total_items_in_acts += 1
            total_value_generated += osum

        rows_processed += 1

    stats = {
        "rows_processed": rows_processed,
        "rows_skipped": rows_skipped,
        "owners_skipped": owners_skipped,
        "total_items_in_acts": total_items_in_acts,
        "total_value_generated": total_value_generated,
    }
    return per_owner, stats


# ---------------------- Helpers: numbers & money -> words ----------------------

def _normalize_apostrophe(s: str) -> str:
    # prefer typographic apostrophe used in Ukrainian texts
    return s.replace("'", "’")


def int_to_words(n: int, lang: str = "uk") -> str:
    """Return cardinal words for integer `n` in specified language."""
    try:
        if n == 0:
            return _normalize_apostrophe(num2words(0, lang=lang))
        w = num2words(n, lang=lang)
        return _normalize_apostrophe(w)
    except Exception:
        # fallback to str
        return str(n)


def money_to_words(amount: Decimal, lang: str = "uk") -> str:
    """Convert Decimal amount -> Ukrainian words with `грн.` and `коп.`

    Examples:
        Decimal('55494.00') -> "п’ятдесят п’ять тисяч чотириста дев’яносто чотири грн. 00 коп."
    """
    q = amount.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    total_kop = int((q * 100).to_integral_value(rounding=ROUND_HALF_UP))
    hryv = total_kop // 100
    kop = total_kop % 100
    hryv_words = int_to_words(hryv, lang=lang)
    # Ensure two-digit fractional part
    return f"{hryv_words} грн. {kop:02d} коп."


def build_mapping_for_owner(data: Dict[str, Any], dept: Dict[str, str]) -> Dict[str, str]:
    """Return mapping for placeholders (strings) including numeric and words for qty and sum."""
    tot_qty = int(data.get("tot_qty", 0))
    tot_sum = data.get("tot_sum", Decimal('0.00'))

    mapping = {
        "TotalQuantityWords": int_to_words(tot_qty, lang='uk'),
        "TotalQuantityNumeric": str(tot_qty),
        "TotalSumNumeric": fmt_number(tot_sum),
        "TotalSumWords": money_to_words(tot_sum, lang='uk'),
        "SecondDirectorPosition": dept.get("position", ""),
        "SecondDirectorName": dept.get("normalized", ""),
        "Val": fmt_number(tot_sum),
    }
    return mapping


def safe_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", name)



def clone_row_preserve_fonts(row):
    """Deep-copy a table row preserving run fonts (v03_force_font_run)."""
    new_row = copy.deepcopy(row._tr)
    return new_row

def set_cell_text_with_font(cell, text, font_name="Times New Roman", font_size=12):
    """Replace cell text while preserving font."""
    cell.text = ""
    p = cell.add_paragraph()
    run = p.add_run(str(text))
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(font_size)


def insert_items_v03_force_font(asset_table, header_idx, items):
    fmt_row = asset_table.rows[header_idx + 1] if header_idx + 1 < len(asset_table.rows) else asset_table.rows[header_idx]
    fmt_tr = fmt_row._tr

    # Determine insertion point (last row before numbering row)
    insert_after_tr = fmt_tr
    for r in asset_table.rows[header_idx + 1:]:
        first_cell_text = (r.cells[0].text or "").strip()
        if first_cell_text == "1":  # numbering row
            break
        insert_after_tr = r._tr

    cursor_tr = insert_after_tr
    for it in items:
        clone = copy.deepcopy(fmt_tr)
        cursor_tr.addnext(clone)
        cursor_tr = clone
        tgt_row = asset_table.rows[[r._tr for r in asset_table.rows].index(clone)]

        unit_price_formatted = fmt_number(it.get("unit_price", Decimal("0.00"))) if it.get("unit_price") is not None else ""
        sum_formatted = fmt_number(it.get("sum", Decimal("0.00"))) if it.get("sum") is not None else ""

        values = [
            str(it.get("name", "")),
            str(it.get("inventory", "")),
            str(it.get("unit", "")),
            str(int(it.get("qty", 0))),
            unit_price_formatted,
            sum_formatted,
            str(it.get("note", "")),
        ]

        for idx, (tgt_cell, val) in enumerate(zip(tgt_row.cells, values)):
            # Preserve existing runs but clear text
            if tgt_cell.paragraphs:
                p = tgt_cell.paragraphs[0]
                p.clear()  # remove existing text but keep paragraph properties
                run = p.add_run(str(val))
            else:
                p = tgt_cell.add_paragraph()
                run = p.add_run(str(val))

            # Font settings
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(9)

            # Alignment
            if idx != 0:  # center all columns except first
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def replace_placeholder_preserve_runs(paragraph, mapping: dict):
    """Replace placeholders in a paragraph while preserving run formatting."""
    for k, v in mapping.items():
        placeholder = f"%{k}%"
        # Iterate over runs carefully
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(v))

def replace_placeholders_doc(doc: Document, mapping: dict):
    """Replace all placeholders in the document while preserving formatting."""
    # Paragraphs
    for p in doc.paragraphs:
        replace_placeholder_preserve_runs(p, mapping)
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholder_preserve_runs(p, mapping)

def save_docx_locally(template_path: str, output_path: str, mapping: dict, items: list):
    doc = Document(template_path)

    replace_placeholders_doc(doc, mapping)

    asset_table, header_idx = None, None
    for t in doc.tables:
        for i, row in enumerate(t.rows):
            if any("Назва об’єкта" in (c.text or "") for c in row.cells):
                asset_table, header_idx = t, i
                break
        if asset_table:
            break
    if not asset_table:
        raise ValueError("Asset table not found")

    insert_items_v03_force_font(asset_table, header_idx, items)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

def create_act_docs(drive_service, docs_service, per_owner: Dict[str, Any], use_local_fallback: bool) -> List[Dict[str, Any]]:
    created = []
    for code, data in per_owner.items():
        if not data["items"]:
            info(f"Owner {code} has no items; skipping.")
            continue

        dept = data["dept"]

        file_name = FILE_NAME_PATTERN.format(
            deptname=dept.get("code"),
        )

        mapping = build_mapping_for_owner(data, dept)

        if use_local_fallback:
            try:
                out_path = os.path.join(OUTPUT_LOCAL_DIR, f"{safe_filename(file_name)}.docx")

                save_docx_locally(
                    template_path="template.docx",
                    output_path=out_path,
                    mapping=mapping,
                    items=data["items"]
                )
                info(f'Created local doc "{out_path}" - items={len(data["items"])} - sum={fmt_number(data["tot_sum"]) }')
                created.append({
                    "local_path": out_path,
                    "name": file_name,
                    "items": len(data["items"]),
                    "sum": data["tot_sum"]
                })
            except Exception as e:
                error(f"Local save failed for {code}: {e}")
            continue

        # Drive-based creation (unchanged path)
        try:
            copy_body = {"name": file_name, "parents": [OUTPUT_FOLDER_ID]}
            newfile = drive_service.files().copy(fileId=TEMPLATE_DOC_ID, body=copy_body, fields="id,name").execute()
            new_id = newfile["id"]
        except HttpError as e:
            error(f"Failed to copy template for {code}: {e}")
            continue

        # Refresh mapping if we prefer fullname in docs
        mapping = build_mapping_for_owner(data, dept)
        # ensure SecondDirectorName is fullname here
        mapping["SecondDirectorName"] = dept.get("fullname", mapping.get("SecondDirectorName"))

        try:
            # replace placeholders and append simple text table
            requests = []
            for key, val in mapping.items():
                token = f"%{key}%"
                requests.append({
                    "replaceAllText": {
                        "containsText": {"text": token, "matchCase": True},
                        "replaceText": val
                    }
                })
            if requests:
                docs_service.documents().batchUpdate(documentId=new_id, body={"requests": requests}).execute()
            # append text block
            lines = []
            hdr = ["№", "Назва", "Інв. номер", "Од.", "К-сть", "Ціна од.", "Сума", "Примітка"]
            lines.append(" | ".join(hdr))
            lines.append("-" * 80)
            for i, it in enumerate(data["items"], start=1):
                row = [
                    str(i), it["name"], str(it["inventory"]), str(it["unit"]),
                    str(it["qty"]), fmt_number(it["unit_price"]), fmt_number(it["sum"]), it.get("note", "")
                ]
                lines.append(" | ".join(row))
            lines.append("-" * 80)
            lines.append(f"Всього: К-сть = {data['tot_qty']}; Сума = {fmt_number(data['tot_sum'])}")
            block = "\n".join(lines) + "\n\n"
            doc = docs_service.documents().get(documentId=new_id).execute()
            end_index = doc.get("body", {}).get("content", [])[-1].get("endIndex", 1)
            docs_service.documents().batchUpdate(documentId=new_id, body={
                "requests": [
                    {"insertText": {"location": {"index": end_index - 1}, "text": block}}
                ]
            }).execute()
            info(f'Created doc "{file_name}" - drive_file_id={new_id} - items={len(data["items"])} - sum={fmt_number(data["tot_sum"]) }')
            created.append({"file_id": new_id, "name": file_name, "items": len(data["items"]), "sum": data["tot_sum"]})
        except HttpError as e:
            error(f"Docs API error while filling doc for {code}: {e}")
            try:
                drive_service.files().delete(fileId=new_id).execute()
            except Exception:
                pass
            continue

    return created


def check_drive_permissions(drive_service) -> Tuple[bool, bool]:
    # returns (drive_ok, use_local_fallback)
    try:
        drive_service.files().get(fileId=TEMPLATE_DOC_ID, fields="id").execute()
    except HttpError as e:
        error(f"Cannot access template doc {TEMPLATE_DOC_ID}: {e}")
        # still allow local fallback if docs API is accessible; but if docs access fails later, script will error.
        return False, True

    if not OUTPUT_FOLDER_ID or OUTPUT_FOLDER_ID.startswith("FOLDER_ID_HERE"):
        warning("OUTPUT_FOLDER_ID not set; using local output only.")
        return False, True

    try:
        drive_service.files().get(fileId=OUTPUT_FOLDER_ID, fields="id").execute()
    except HttpError as e:
        # detect storage quota exceeded
        msg = str(e)
        if "storageQuotaExceeded" in msg or "quota" in msg:
            warning("Drive storage quota exceeded; will save output locally to output_docs/.")
            return False, True
        error(f"Cannot access output folder {OUTPUT_FOLDER_ID}: {e}")
        return False, True

    # test create & delete small file
    test_body = {"name": "asset_acts_perm_test.tmp", "parents": [OUTPUT_FOLDER_ID], "mimeType": "application/vnd.google-apps.document"}
    try:
        created = drive_service.files().create(body=test_body, fields="id").execute()
        drive_service.files().delete(fileId=created["id"]).execute()
    except HttpError as e:
        msg = str(e)
        if "storageQuotaExceeded" in msg or "quota" in msg:
            warning("Drive storage quota exceeded; will save output locally to output_docs/.")
            return False, True
        error(f"Service account cannot create files in output folder: {e}")
        return False, True

    info("Drive template and folder accessible.")
    return True, False


def main():
    check_constants()
    sheets_svc, drive_svc, docs_svc = build_services()
    
        # verify that the provided ASSETS and DEPARTMENTS IDs are actual Google Spreadsheets
    try:
        ensure_file_is_spreadsheet(drive_svc, ASSETS_SPREADSHEET_ID, "Assets spreadsheet")
        ensure_file_is_spreadsheet(drive_svc, DEPARTMENTS_SPREADSHEET_ID, "Departments spreadsheet")
    except SystemExit:
        raise
    
    drive_ok, use_local = check_drive_permissions(drive_svc)
    
    
    try:
        # ensure template contains placeholders (best-effort)
        try:
            doc = docs_svc.documents().get(documentId=TEMPLATE_DOC_ID).execute()
            txt = ""
            for el in doc.get("body", {}).get("content", []):
                for e in el.get("paragraph", {}).get("elements", []):
                    txt += e.get("textRun", {}).get("content", "")
            if "%" not in txt and TABLE_PLACEHOLDER_TOKEN not in txt:
                warning("Template does not contain %...% placeholders nor table placeholder; script will still proceed.")
        except HttpError as ex:
            warning(f"Could not read template placeholders (Docs API): {ex}; proceeding (local fallback may use basic template).")
    except Exception:
        pass

    info(f"Start processing assets spreadsheet (ID={ASSETS_SPREADSHEET_ID})")
    departments = load_departments(sheets_svc)
    per_owner, stats = parse_assets(sheets_svc, departments)

    if not per_owner:
        summary("No valid owners/items found; nothing to generate.")
        return

    created = create_act_docs(drive_svc if drive_ok else None, docs_svc, per_owner, use_local or not drive_ok)

    summary(f"rows_processed={stats['rows_processed']}, rows_skipped={stats['rows_skipped']}, owners_skipped={stats['owners_skipped']}, acts_generated={len(created)}, items_in_acts={stats['total_items_in_acts']}, total_value_generated={fmt_number(stats['total_value_generated'])}")

    _logfile.close()


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        error(f"Unhandled exception: {exc}")
        _logfile.close()
        sys.exit(1)
