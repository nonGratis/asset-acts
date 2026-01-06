import copy
from decimal import Decimal
from typing import Dict, Any

from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
from .formatters import fmt_number, money_to_words


def build_mapping_for_owner(data: Dict[str, Any], dept: Dict[str, str]) -> Dict[str, str]:
    """
    Args:
        data: Owner data containing items, quantities, and sums
        dept: Department information dictionary

    Returns:
        Dictionary mapping placeholder names to their values
    """
    tot_qty = int(data.get("tot_qty", 0))
    tot_sum = data.get("tot_sum", Decimal("0.00"))

    mapping = {
        "TotalQuantityWords": num2words(tot_qty, lang="uk"),
        "TotalQuantityNumeric": str(tot_qty),
        "TotalSumNumeric": fmt_number(tot_sum),
        "TotalSumWords": money_to_words(tot_sum, lang="uk"),
        "SecondDirectorPosition": dept.get("position", ""),
        "SecondDirectorName": dept.get("normalized", ""),
        "Val": fmt_number(tot_sum),
    }
    return mapping


def replace_placeholder_preserve_runs(paragraph, mapping: dict):
    """Replace placeholders in a paragraph while preserving run formatting"""
    for k, v in mapping.items():
        placeholder = f"%{k}%"
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(v))


def replace_placeholders_doc(doc: DocumentType, mapping: dict):
    """Replace all placeholders in the document while preserving formatting"""
    for p in doc.paragraphs:
        replace_placeholder_preserve_runs(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholder_preserve_runs(p, mapping)


def add_assets_preserve_formatting(asset_table, header_idx, items):
    """
    Args:
        asset_table: docx Table object containing asset data
        header_idx: Index of the header row
        items: List of asset item dictionaries to add
    """
    fmt_row = (
        asset_table.rows[header_idx + 1]
        if header_idx + 1 < len(asset_table.rows)
        else asset_table.rows[header_idx]
    )
    fmt_tr = fmt_row._tr

    # Find totals row (containing 'всього' or starting with '1')
    totals_row = None
    for r in asset_table.rows[header_idx + 1 :]:
        first = (r.cells[0].text or "").strip().lower()
        if first.startswith("всього"):
            totals_row = r
            break

    if totals_row is None:
        for r in asset_table.rows[header_idx + 1 :]:
            if (r.cells[0].text or "").strip() == "1":
                totals_row = r
                break

    # If still not found, insert before last row
    if totals_row is None:
        totals_tr = asset_table.rows[-1]._tr
    else:
        totals_tr = totals_row._tr

    # Insert items before totals row
    for it in items:
        clone_tr = copy.deepcopy(fmt_tr)
        totals_tr.addprevious(clone_tr)

        # Find the newly inserted row object
        new_row = next(r for r in asset_table.rows if r._tr == clone_tr)

        unit_price_formatted = (
            fmt_number(it.get("unit_price", Decimal("0.00")))
            if it.get("unit_price") is not None
            else ""
        )
        sum_formatted = (
            fmt_number(it.get("sum", Decimal("0.00"))) if it.get("sum") is not None else ""
        )

        values = [
            str(it.get("name", "")),
            str(it.get("inventory", "")),
            str(it.get("unit", "")),
            str(int(it.get("qty", 0))),
            unit_price_formatted,
            sum_formatted,
            str(it.get("note", "")),
        ]
        
        for idx, (tgt_cell, val) in enumerate(zip(new_row.cells, values)):
            if tgt_cell.paragraphs: # Write into cells preserving font settings
                p = tgt_cell.paragraphs[0]
                try:
                    p.clear()
                except Exception:
                    for run in list(p.runs):
                        run.text = ""
            else:
                p = tgt_cell.add_paragraph()

            run = p.add_run(str(val))

            run.font.name = "Times New Roman"
            try:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            except Exception:
                pass
            run.font.size = Pt(9)

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 0 else WD_ALIGN_PARAGRAPH.LEFT
