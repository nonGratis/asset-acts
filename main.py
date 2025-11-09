import os
import re
import sys
import copy
import logging
from io import BytesIO

import fitz
from PIL import Image
from dotenv import load_dotenv

from docx import Document
from docx.document import Document as DocumentType
from decimal import Decimal, ROUND_HALF_UP
from typing import List, Dict, Any, Tuple

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload

from num2words import num2words
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert as docx2pdf_convert

load_dotenv()

SERVICE_ACCOUNT_KEYFILE = os.getenv("GOOGLE_CREDS_PATH", "credentials.json")

ASSETS_SPREADSHEET_ID = os.getenv("ASSETS_SHEET_ID", "")
ASSETS_SHEET_NAME = "list"

DEPARTMENTS_SPREADSHEET_ID = os.getenv("DEPARTMENTS_SHEET_ID", "")
DEPARTMENTS_SHEET_NAME = "Department"

SHARED_DRIVE_ID = os.getenv("SHARED_DRIVE_ID", "")

OUTPUT_LOCAL_DIR_DOC = "docs"
OUTPUT_LOCAL_DIR_PDF = "pdfs"

# Columns (1-based)
COL_ID = 1
COL_NAME = 2
COL_INVENTORY_NUMBER = 3
COL_UNIT = 4
COL_QUANTITY = 5
COL_PRICE = 6
COL_OWNERS = 7
COL_DATE = 8
COL_GENERATE_FLAG = 9

DEPT_COL_CODE = 1
DEPT_COL_TYPE = 2
DEPT_COL_STATUS = 3
DEPT_COL_POSITION = 4
DEPT_COL_FULLNAME = 5
DEPT_COL_NORMALIZED = 6

FILE_NAME_PATTERN = "Акт. {deptname}"  
THOUSAND_SEPARATOR = " "
DECIMAL_SEPARATOR = ","
CURRENCY_SUFFIX = ""
ALLOW_ROUNDING_ADJUST = True

logging.basicConfig(
    level=logging.INFO,
    format="%(levelname)s: %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
log = logging.getLogger(__name__)

def check_constants() -> None:
    missing = []
    if not os.path.isfile(SERVICE_ACCOUNT_KEYFILE):
        missing.append(f"SERVICE_ACCOUNT_KEYFILE file not found: {SERVICE_ACCOUNT_KEYFILE}")
    if not os.path.isfile("template.docx"):
        missing.append("Template file not found: template.docx")
    for keyname, value in (
        ("ASSETS_SPREADSHEET_ID", ASSETS_SPREADSHEET_ID),
        ("DEPARTMENTS_SPREADSHEET_ID", DEPARTMENTS_SPREADSHEET_ID),
        ("SHARED_DRIVE_ID", SHARED_DRIVE_ID),
    ):
        if not value:
            missing.append(f"Missing constant {keyname}")
    if missing:
        for m in missing:
            log.error(m)
        raise SystemExit("Missing critical constants; aborting.")


SCOPES = [
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/documents",
    "https://www.googleapis.com/auth/spreadsheets.readonly",
]


def build_services():
    creds = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_KEYFILE, scopes=SCOPES)
    sheets = build("sheets", "v4", credentials=creds, cache_discovery=False)
    drive = build("drive", "v3", credentials=creds, cache_discovery=False)
    docs = build("docs", "v1", credentials=creds, cache_discovery=False)
    return sheets, drive, docs

def ensure_file_is_spreadsheet(drive_service, file_id: str, label: str) -> None:
    """
    Verify that `file_id` exists and is a Google Spreadsheet.
    If not: log ERROR and exit.
    """
    try:
        meta = drive_service.files().get(fileId=file_id, fields="id, name, mimeType").execute()
    except HttpError as e:
        log.error(f"Cannot fetch {label} (id={file_id}): {e}")
        raise SystemExit(1)
    mime = meta.get("mimeType", "")
    if mime != "application/vnd.google-apps.spreadsheet":
        log.error(f"{label} (id={file_id}) is not a Google Spreadsheet (mimeType={mime}). Please provide the correct spreadsheet ID.")
        raise SystemExit(1)
    log.info(f"{label} found: {meta.get('name', '<untitled>')} (id={file_id})")


def safe_get(row: list, col: int, default=""):
    """Return 1-based column from row safely.

    row: list of cell values as returned by Sheets API
    col: 1-based column index
    """
    if row is None:
        return default
    idx = col - 1
    if idx < 0:
        return default
    if idx >= len(row):
        return default
    val = row[idx]
    return val if val is not None else default


def parse_number(s) -> Decimal:
    if s is None:
        log.error("Empty numeric")
        raise ValueError("Empty numeric")
    st = str(s).strip()
    st = st.replace("\xa0", "").replace(" ", "")
    st = st.replace(",", ".")
    if st == "":
        log.error("Empty numeric")
        raise ValueError("Empty numeric")
    try:
        return Decimal(st)
    except Exception as e:
        log.error("Invalid numeric '%s': %s", s, e)
        raise


def quantize_money(d: Decimal) -> Decimal:
    return d.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def fmt_number(val: Decimal) -> str:
    q = quantize_money(Decimal(val))
    s = f"{q:,.2f}"
    s = s.replace(",", "TEMP_THOUS").replace(".", "TEMP_DEC")
    s = s.replace("TEMP_THOUS", THOUSAND_SEPARATOR).replace("TEMP_DEC", DECIMAL_SEPARATOR)
    return s + (CURRENCY_SUFFIX or "")


def normalize_code(token: str) -> str:
    return re.sub(r"\s+", "", token).upper()


def read_sheet_values(sheets_service, spreadsheet_id: str, sheet_name: str):
    rng = f"{sheet_name}"
    try:
        res = sheets_service.spreadsheets().values().get(spreadsheetId=spreadsheet_id, range=rng).execute()
        return res.get("values", [])
    except HttpError as e:
        msg = str(e)
        if "This operation is not supported for this document" in msg or "not supported for this document" in msg:
            log.error(f"Sheets API: file id {spreadsheet_id} is not a spreadsheet or not supported by Sheets API. Check that the ID points to a Google Sheet and the service account has access.")
            raise SystemExit(1)
        log.error(f"Sheets API error for spreadsheet={spreadsheet_id}, range={rng}: {e}")
        raise


def load_departments(sheets_service) -> Dict[str, Dict[str, str]]:
    vals = read_sheet_values(sheets_service, DEPARTMENTS_SPREADSHEET_ID, DEPARTMENTS_SHEET_NAME)
    depts = {}
    if not vals or len(vals) < 2:
        log.warning("Departments sheet empty or missing rows.")
        return depts
    for i, row in enumerate(vals[1:], start=2):
        code = str(safe_get(row, DEPT_COL_CODE, "")).strip()
        if not code:
            log.warning(f"Departments row {i} missing code; skipping.")
            continue
        key = normalize_code(code)
        depts[key] = {
            "code": safe_get(row, DEPT_COL_CODE, ""),
            "type": safe_get(row, DEPT_COL_TYPE, ""),
            "status": safe_get(row, DEPT_COL_STATUS, ""),
            "position": safe_get(row, DEPT_COL_POSITION, ""),
            "fullname": safe_get(row, DEPT_COL_FULLNAME, ""),
            "normalized": safe_get(row, DEPT_COL_NORMALIZED, ""),
        }
    return depts


def parse_owner_token(tok: str) -> Tuple[str, int, bool]:
    tok = tok.strip()
    m = re.match(r"^(.*?)-\s*([0-9]+)\s*$", tok)
    if m:
        return m.group(1).strip(), int(m.group(2)), True
    return tok, None, False


def parse_assets(sheets_service, departments: Dict[str, Dict[str, str]]):
    vals = read_sheet_values(sheets_service, ASSETS_SPREADSHEET_ID, ASSETS_SHEET_NAME)
    if not vals or len(vals) < 2:
        log.info("No assets rows found.")
        return {}, {"rows_processed": 0, "rows_skipped": 0, "owners_skipped": 0, "total_items_in_acts": 0, "total_value_generated": Decimal("0.00")}

    rows_processed = 0
    rows_skipped = 0
    owners_skipped = 0
    per_owner: Dict[str, Dict[str, Any]] = {}
    total_items_in_acts = 0
    total_value_generated = Decimal("0.00")

    for rindex, row in enumerate(vals[1:], start=2):
        if not any(str(cell).strip() for cell in row):
            continue

        gen_flag = safe_get(row, COL_GENERATE_FLAG, "")
        if str(gen_flag).strip().upper() != "TRUE":
            rows_skipped += 1
            continue

        try:
            name = safe_get(row, COL_NAME, "")
            invnum = safe_get(row, COL_INVENTORY_NUMBER, "")
            unit = safe_get(row, COL_UNIT, "").lower()
            qty_raw = safe_get(row, COL_QUANTITY, "")
            price_raw = safe_get(row, COL_PRICE, "")
            owners_raw = safe_get(row, COL_OWNERS, "")

            if not name:
                log.warning(f"Row {rindex} missing name; skipping.")
                rows_skipped += 1
                continue

            qty = int(parse_number(qty_raw))
            price = parse_number(price_raw)

            if qty <= 0:
                log.error(f"Row {rindex} has non-positive quantity {qty}; skip row.")
                rows_skipped += 1
                continue

            unit_price = quantize_money(price / Decimal(qty))
        except Exception as e:
            log.error(f"Row {rindex} parse error: {e}; skipping row.")
            rows_skipped += 1
            continue

        tokens = [t.strip() for t in str(owners_raw).split(",") if t.strip()]
        if not tokens:
            log.error(f"Row {rindex} no owners; skip.")
            rows_skipped += 1
            continue

        token_infos = []
        any_explicit = False
        for tok in tokens:
            base, num, explicit = parse_owner_token(tok)
            if explicit:
                any_explicit = True
            token_infos.append((base, num, explicit))

        if any_explicit:
            if not all(t[2] for t in token_infos):
                log.error(f"Row {rindex} mixed explicit and implicit owners; skip.")
                rows_skipped += 1
                continue
            total_spec = sum(t[1] for t in token_infos)
            if total_spec != qty:
                log.error(f"Row {rindex} owner counts sum {total_spec} != quantity {qty}; skip.")
                rows_skipped += 1
                continue
        else:
            if len(token_infos) != 1:
                log.error(f"Row {rindex} ambiguous multiple owners without counts; skip.")
                rows_skipped += 1
                continue
            base = token_infos[0][0]
            token_infos[0] = (base, qty, True)

        owners_for_row = []
        for base, num, _ in token_infos:
            key = normalize_code(base)
            dept = departments.get(key)
            if not dept:
                log.error(f"Row {rindex} owner '{base}' not found; skipping this owner entry.")
                owners_skipped += 1
                continue
            owners_for_row.append((key, int(num), dept))

        if not owners_for_row:
            log.info(f"Row {rindex}: all owners skipped; skipping row.")
            rows_skipped += 1
            continue

        owner_sums = []
        for (_, oqty, _) in owners_for_row:
            owner_sums.append(quantize_money(unit_price * Decimal(oqty)))
        sum_owner_sums = sum(owner_sums)
        price_q = quantize_money(price)
        if sum_owner_sums != price_q:
            diff = price_q - sum_owner_sums
            if ALLOW_ROUNDING_ADJUST:
                owner_sums[-1] = quantize_money(owner_sums[-1] + diff)
                log.warning(f"Row {rindex} rounding adjusted by {diff} on last owner.")
            else:
                log.warning(f"Row {rindex} rounding mismatch {price_q - sum_owner_sums}; continuing.")

        for (key, oqty, dept), osum in zip(owners_for_row, owner_sums):
            if key not in per_owner:
                per_owner[key] = {"dept": dept, "items": [], "tot_qty": 0, "tot_sum": Decimal("0.00")}
            per_owner[key]["items"].append({
                "name": name,
                "inventory": invnum,
                "unit": unit,
                "qty": int(oqty),
                "unit_price": unit_price,
                "sum": osum,
                "note": "",
            })
            per_owner[key]["tot_qty"] += int(oqty)
            per_owner[key]["tot_sum"] += osum
            total_items_in_acts += 1
            total_value_generated += osum

        rows_processed += 1

    stats = {
        "rows_processed": rows_processed,
        "rows_skipped": rows_skipped,
        "owners_skipped": owners_skipped,
        "total_items_in_acts": total_items_in_acts,
        "total_value_generated": total_value_generated,
    }
    return per_owner, stats


def money_to_words(amount: Decimal, lang: str = "uk") -> str:
    q = quantize_money(amount)
    total_kop = int((q * 100).to_integral_value(rounding=ROUND_HALF_UP))
    hryv = total_kop // 100
    kop = total_kop % 100

    def _form_for(n: int, forms: tuple) -> str:
        n = abs(int(n))
        if n % 10 == 1 and n % 100 != 11:
            return forms[0]
        if n % 10 in (2, 3, 4) and n % 100 not in (12, 13, 14):
            return forms[1]
        return forms[2]

    if hryv == 0:
        hryv_words = "нуль"
    else:
        thousands = hryv // 1000
        rest = hryv % 1000
        parts = []
        if thousands:
            parts.append(num2words(thousands, lang=lang))
            parts.append(_form_for(thousands, ("тисяча", "тисячі", "тисяч")))
        if rest:
            parts.append(num2words(rest, lang=lang))
        hryv_words = " ".join(parts)

    hryv_words = re.sub(r'\bодин\b', 'одна', hryv_words)
    hryv_words = re.sub(r'\bдва\b', 'дві', hryv_words) 

    return f"{hryv_words} грн. {kop:02d} коп."


def build_mapping_for_owner(data: Dict[str, Any], dept: Dict[str, str]) -> Dict[str, str]:
    """Return mapping for placeholders (strings) including numeric and words for qty and sum."""
    tot_qty = int(data.get("tot_qty", 0))
    tot_sum = data.get("tot_sum", Decimal('0.00'))

    mapping = {
        "TotalQuantityWords": num2words(tot_qty, lang='uk'),
        "TotalQuantityNumeric": str(tot_qty),
        "TotalSumNumeric": fmt_number(tot_sum),
        "TotalSumWords": money_to_words(tot_sum, lang='uk'),
        "SecondDirectorPosition": dept.get("position", ""),
        "SecondDirectorName": dept.get("normalized", ""),
        "Val": fmt_number(tot_sum),
    }
    return mapping


def safe_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", name)

def add_assets_preserve_formatting(asset_table, header_idx, items):
    fmt_row = asset_table.rows[header_idx + 1] if header_idx + 1 < len(asset_table.rows) else asset_table.rows[header_idx]
    fmt_tr = fmt_row._tr

    # find totals row first cell containing 'всього'
    totals_row = None
    for r in asset_table.rows[header_idx + 1:]:
        first = (r.cells[0].text or "").strip().lower()
        if first.startswith("всього"):
            totals_row = r
            break

    # find numbering row where first cell == "1"
    if totals_row is None:
        for r in asset_table.rows[header_idx + 1:]:
            if (r.cells[0].text or "").strip() == "1":
                totals_row = r
                break

    # If still not, append before the last row
    if totals_row is None:
        totals_tr = asset_table.rows[-1]._tr
    else:
        totals_tr = totals_row._tr

    # insrt items before totals_tr
    for it in items:
        clone_tr = copy.deepcopy(fmt_tr)
        totals_tr.addprevious(clone_tr)

        # find the newly inserted row object in asset_table.rows
        new_row = next(r for r in asset_table.rows if r._tr == clone_tr)

        unit_price_formatted = fmt_number(it.get("unit_price", Decimal("0.00"))) if it.get("unit_price") is not None else ""
        sum_formatted = fmt_number(it.get("sum", Decimal("0.00"))) if it.get("sum") is not None else ""

        values = [
            str(it.get("name", "")),
            str(it.get("inventory", "")),
            str(it.get("unit", "")),
            str(int(it.get("qty", 0))),
            unit_price_formatted,
            sum_formatted,
            str(it.get("note", "")),
        ]

        # write into cells preserving font settings
        for idx, (tgt_cell, val) in enumerate(zip(new_row.cells, values)):
            if tgt_cell.paragraphs:
                p = tgt_cell.paragraphs[0]
                try:
                    p.clear()  # remove existing text but keep paragraph properties
                except Exception:
                    for run in list(p.runs):
                        run.text = ""
            else:
                p = tgt_cell.add_paragraph()

            run = p.add_run(str(val))

            run.font.name = 'Times New Roman'
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            except Exception:
                pass
            run.font.size = Pt(9)

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 0 else WD_ALIGN_PARAGRAPH.LEFT

def replace_placeholder_preserve_runs(paragraph, mapping: dict):
    """Replace placeholders in a paragraph while preserving run formatting."""
    for k, v in mapping.items():
        placeholder = f"%{k}%"
        # Iterate over runs carefully
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(v))

def replace_placeholders_doc(doc: DocumentType, mapping: dict):
    """Replace all placeholders in the document while preserving formatting."""
    # Paragraphs
    for p in doc.paragraphs:
        replace_placeholder_preserve_runs(p, mapping)
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholder_preserve_runs(p, mapping)

def save_docx_locally(template_path: str, output_path: str, mapping: dict, items: list):
    doc = Document(template_path)

    replace_placeholders_doc(doc, mapping)

    asset_table, header_idx = None, None
    for t in doc.tables:
        for i, row in enumerate(t.rows):
            if any("Назва об’єкта" in (c.text or "") for c in row.cells):
                asset_table, header_idx = t, i
                break
        if asset_table:
            break
    if not asset_table:
        raise ValueError("Asset table not found")

    add_assets_preserve_formatting(asset_table, header_idx, items)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

def upload_to_drive(drive_service, file_path: str, file_name: str) -> str:
    """Upload file to Google Drive shared drive."""
    try:
        media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        file_metadata = {
            'name': file_name,
            'parents': [SHARED_DRIVE_ID]
        }
        
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            supportsAllDrives=True
        ).execute()
        
        return file.get('id')
        
    except HttpError as e:
        raise RuntimeError(f"Drive upload failed: {e}")

def convert_to_pdf(docx_path: str) -> str:
    """Convert DOCX to PDF. Returns PDF path on success."""

    pdf_path = docx_path.replace(".docx", ".pdf")
    pdf_dir = os.path.dirname(pdf_path).replace(OUTPUT_LOCAL_DIR_DOC, OUTPUT_LOCAL_DIR_PDF)
    pdf_path = os.path.join(pdf_dir, os.path.basename(pdf_path))
    
    os.makedirs(pdf_dir, exist_ok=True)
    
    try:
        docx2pdf_convert(docx_path, pdf_path)
        return pdf_path
    except Exception as e:
        raise RuntimeError(f"PDF conversion failed: {e}")

def create_act_docs_local(per_owner: Dict[str, Any], drive_service) -> List[Dict[str, Any]]:
    """Save all documents locally and upload to Google Drive"""
    created = []
    upload_failed = []
    
    for code, data in per_owner.items():
        if not data["items"]:
            log.info(f"Owner {code} has no items; skipping.")
            continue

        dept = data["dept"]
        file_name = FILE_NAME_PATTERN.format(deptname=dept.get("code"))
        mapping = build_mapping_for_owner(data, dept)

        try:
            # Create DOCX
            docx_path = os.path.join(OUTPUT_LOCAL_DIR_DOC, f"{safe_filename(file_name)}.docx")
            save_docx_locally(
                template_path="template.docx",
                output_path=docx_path,
                mapping=mapping,
                items=data["items"],
            )
            
            doc_info = {
                "docx_path": docx_path,
                "name": file_name,
                "items": len(data["items"]),
                "sum": data["tot_sum"],
            }
            
            # Upload to Google Drive
            try:
                drive_file_id = upload_to_drive(drive_service, docx_path, f"{file_name}.docx")
                doc_info["drive_file_id"] = drive_file_id
                log.info(f'Created and uploaded "{file_name}.docx" (ID: {drive_file_id}) - items={len(data["items"])} - sum={fmt_number(data["tot_sum"])}')
            except Exception as e:
                log.warning(f"Drive upload failed for {code}: {e}")
                upload_failed.append(code)
                log.info(f'Created local "{docx_path}" (upload failed) - items={len(data["items"])} - sum={fmt_number(data["tot_sum"])}')
            
            created.append(doc_info)
                        
            # Create PDF and JPEG
            try:
                pdf_path = convert_to_pdf(docx_path)
                doc_info["pdf_path"] = pdf_path
                
                try:
                    jpeg_path = convert_to_jpeg(pdf_path)
                    doc_info["jpeg_path"] = jpeg_path
                    log.info(f'Created docs "{docx_path}" + PDF + JPEG - items={len(data["items"])} - sum={fmt_number(data["tot_sum"])}')
                except Exception as jpeg_err:
                    log.warning(f"JPEG conversion failed for {code}: {jpeg_err}")
                    log.info(f'Created docs "{docx_path}" + PDF - items={len(data["items"])} - sum={fmt_number(data["tot_sum"])}')
                
            except Exception as e:
                log.warning(f"PDF conversion failed for {code}: {e}")
                log.info(f'Created doc "{docx_path}" (PDF/JPEG skipped) - items={len(data["items"])} - sum={fmt_number(data["tot_sum"])}')

            
        except Exception as e:
            log.error(f"Document creation failed for {code}: {e}")
            continue

    if upload_failed:
        log.info(f"Drive upload failed for: {', '.join(upload_failed)}")

    return created

def convert_to_jpeg(pdf_path: str) -> str:   
    jpeg_path = pdf_path.replace(".pdf", ".jpg")
    jpeg_dir = os.path.dirname(jpeg_path).replace(OUTPUT_LOCAL_DIR_PDF, "jpegs")
    jpeg_path = os.path.join(jpeg_dir, os.path.basename(jpeg_path))
    
    os.makedirs(jpeg_dir, exist_ok=True)
    
    try:
        doc = fitz.open(pdf_path)
        page = doc[0]
        
        # Calculate longest side
        rect = page.rect
        zoom_x = 1280 / max(rect.width, rect.height)
        mat = fitz.Matrix(zoom_x, zoom_x)
        
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_data = pix.tobytes("ppm")
        
        img = Image.open(BytesIO(img_data))
        img = resize_to_max_dimension(img, 1280)
        
        img.save(jpeg_path, "JPEG", quality=100, optimize=False)
        doc.close()
        
        return jpeg_path
        
    except Exception as e:
        raise RuntimeError(f"JPEG conversion failed: {e}")

def resize_to_max_dimension(img: Image.Image, max_size: int) -> Image.Image:
    """Resize image maintaining aspect ratio with max dimension constraint."""
    w, h = img.size
    if max(w, h) <= max_size:
        return img
    
    ratio = max_size / max(w, h)
    new_size = (int(w * ratio), int(h * ratio))
    
    return img.resize(new_size, Image.Resampling.LANCZOS)

def main():
    check_constants()
    sheets_svc, drive_svc, _ = build_services()

    try:
        ensure_file_is_spreadsheet(drive_svc, ASSETS_SPREADSHEET_ID, "Assets spreadsheet")
        ensure_file_is_spreadsheet(drive_svc, DEPARTMENTS_SPREADSHEET_ID, "Departments spreadsheet")
    except SystemExit:
        raise

    log.info(f"Start processing assets spreadsheet (ID={ASSETS_SPREADSHEET_ID})")
    departments = load_departments(sheets_svc)
    per_owner, stats = parse_assets(sheets_svc, departments)

    if not per_owner:
        log.info("No valid owners/items found; nothing to generate.")
        return

    created = create_act_docs_local(per_owner, drive_svc)

    log.info(f"rows_processed={stats['rows_processed']}, rows_skipped={stats['rows_skipped']}, owners_skipped={stats['owners_skipped']}, acts_generated={len(created)}, items_in_acts={stats['total_items_in_acts']}, total_value_generated={fmt_number(stats['total_value_generated'])}")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        log.error(f"Unhandled exception: {exc}")
        sys.exit(1)